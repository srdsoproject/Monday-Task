import datetime
import os
import json
import gspread
import pandas as pd
import smtplib

from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


# -------------------------------------------------
# CONFIG
# -------------------------------------------------
SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
EMAIL_FROM = os.getenv("EMAIL_FROM")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
EMAIL_TO = os.getenv("EMAIL_TO").split(",")

EMAIL_SUBJECT = "Officer Not Done List Report"

EMAIL_BODY = """
Dear user,

Please find attached the Officer not done report.

Regards,
Automation System
"""

# -------------------------------------------------
# DATE / SHEET NAME
# -------------------------------------------------
today = datetime.date.today()
month_name = today.strftime("%B")
year_short = today.strftime("%y")

# Example: April 26
SHEET_NAME = f"{month_name} {year_short}"

# -------------------------------------------------
# GOOGLE CREDENTIALS
# -------------------------------------------------
google_creds_json = os.getenv("GOOGLE_CREDS_JSON")

with open("credentials.json", "w") as f:
    json.dump(json.loads(google_creds_json), f)

# -------------------------------------------------
# REPORT GENERATION
# -------------------------------------------------
def generate_report():

    month_start = today.replace(day=1)
    yesterday = today - datetime.timedelta(days=1)

    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive"
    ]

    creds = ServiceAccountCredentials.from_json_keyfile_name(
        "credentials.json",
        scope
    )

    client = gspread.authorize(creds)

    sheet = client.open_by_key(SPREADSHEET_ID).worksheet(SHEET_NAME)

    data = sheet.get_all_values()

    if not data:
        print("No data found.")
        return None

    df = pd.DataFrame(data[1:], columns=data[0])

    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")

    filtered_df = df[
        (df["Date"] >= pd.to_datetime(month_start)) &
        (df["Date"] < pd.to_datetime(today)) &
        (df["Insp. Done"].str.upper() == "NO") &
        (df["Def. Submitted"].str.upper() == "NO")
    ]

    if filtered_df.empty:
        print("No pending records.")
        return None

    report_df = filtered_df[
        [
            "SN",
            "Date",
            "Name",
            "Designation",
            "Department",
            "Contact",
            "Remarks"
        ]
    ].copy()

    report_df.reset_index(drop=True, inplace=True)
    report_df["Sr. No"] = report_df.index + 1
    report_df["Date"] = report_df["Date"].dt.strftime("%Y-%m-%d")

    report_columns = [
        "Sr. No",
        "Date",
        "Name",
        "Designation",
        "Department",
        "Contact",
        "Remarks"
    ]

    report_df = report_df[report_columns]

    # -------------------------------------------------
    # CREATE WORD REPORT
    # -------------------------------------------------
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    heading_text = (
        f"Officer not done list "
        f"(Period: {month_start:%Y-%m-%d} to {yesterday:%Y-%m-%d})"
    )

    heading = doc.add_heading(heading_text, level=1)

    run = heading.runs[0]
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True
    run.underline = True
    run.font.size = Pt(16)

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(report_columns))
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells

    for i, col_name in enumerate(report_columns):
        hdr_cells[i].text = col_name

        for para in hdr_cells[i].paragraphs:
            for r in para.runs:
                r.bold = True

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rows
    for _, row in report_df.iterrows():

        row_cells = table.add_row().cells

        for i, col_name in enumerate(report_columns):
            row_cells[i].text = str(row[col_name])

            if i in (0, 1):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_filename = f"not_done_report_{today}.docx"

    doc.save(report_filename)

    print("Report created:", report_filename)

    return report_filename


# -------------------------------------------------
# EMAIL
# -------------------------------------------------
def send_email(filename):

    msg = MIMEMultipart()

    msg["From"] = EMAIL_FROM
    msg["To"] = ",".join(EMAIL_TO)
    msg["Subject"] = EMAIL_SUBJECT

    msg.attach(MIMEText(EMAIL_BODY, "plain"))

    with open(filename, "rb") as attachment:

        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    encoders.encode_base64(part)

    part.add_header(
        "Content-Disposition",
        f"attachment; filename={filename}"
    )

    msg.attach(part)

    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(EMAIL_FROM, EMAIL_PASSWORD)
        server.sendmail(
            EMAIL_FROM,
            EMAIL_TO,
            msg.as_string()
        )

    print("Email sent successfully.")


# -------------------------------------------------
# MAIN
# -------------------------------------------------
if __name__ == "__main__":

    # Monday check
    # Monday = 0
    if today.weekday() != 0:
        print("Today is not Monday. Exiting.")
        exit()

    print("Monday confirmed. Running report...")

    file = generate_report()

    if file:
        send_email(file)
